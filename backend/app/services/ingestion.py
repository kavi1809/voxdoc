"""Document text extraction, chunking, and safe URL fetching."""

from __future__ import annotations

import ipaddress
import logging
import socket
from urllib.parse import urlparse

import docx
import httpx
import pandas as pd
import pypdf
from bs4 import BeautifulSoup
from fastapi import HTTPException

from app.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()

MAX_REDIRECTS = 5


# ── Text extraction ────────────────────────────────────────────────────────────


def extract_text_from_pdf(file_path: str) -> str:
    reader = pypdf.PdfReader(file_path)
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def extract_text_from_docx(file_path: str) -> str:
    document = docx.Document(file_path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Tables carry a lot of the meaning in real-world .docx files and were
    # previously dropped entirely.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as fh:
        return fh.read()


# ── SSRF-safe URL fetching ─────────────────────────────────────────────────────


def _assert_public_url(url: str) -> None:
    """
    Reject anything that could reach the internal network.

    Without this the endpoint is a blind SSRF: an attacker could ask the server
    to fetch http://169.254.169.254/ (cloud metadata, often including
    credentials), http://localhost:8000/ (this API), or any host on the private
    network the server sits in — and read the result back as a "document".
    """
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise HTTPException(400, "Only http:// and https:// URLs are supported")
    if not parsed.hostname:
        raise HTTPException(400, "URL has no host")

    try:
        infos = socket.getaddrinfo(parsed.hostname, parsed.port or (443 if parsed.scheme == "https" else 80))
    except socket.gaierror:
        raise HTTPException(400, f"Could not resolve host: {parsed.hostname}")

    for info in infos:
        ip = ipaddress.ip_address(info[4][0])
        if (
            ip.is_private
            or ip.is_loopback
            or ip.is_link_local
            or ip.is_reserved
            or ip.is_multicast
            or ip.is_unspecified
        ):
            raise HTTPException(400, "Refusing to fetch a non-public address")


async def extract_text_from_url(url: str) -> tuple[str, str]:
    """
    Scrape a page and return (title, clean text).

    Redirects are followed manually so that every hop is re-validated — a public
    URL that 302s to 169.254.169.254 would otherwise walk straight past the check.
    """
    current = url
    async with httpx.AsyncClient(follow_redirects=False, timeout=15) as client:
        for _ in range(MAX_REDIRECTS):
            _assert_public_url(current)
            try:
                response = await client.get(current, headers={"User-Agent": "Voxdoc/1.0"})
            except httpx.HTTPError as exc:
                raise HTTPException(400, f"Could not fetch URL: {exc}") from exc

            if response.is_redirect and "location" in response.headers:
                current = str(response.next_request.url)
                continue
            break
        else:
            raise HTTPException(400, "Too many redirects")

    if response.status_code >= 400:
        raise HTTPException(400, f"URL returned HTTP {response.status_code}")

    content = response.content[: settings.max_url_bytes]
    soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
        tag.decompose()

    title = soup.title.string.strip() if soup.title and soup.title.string else url
    text = soup.get_text(separator="\n", strip=True)
    if not text.strip():
        raise HTTPException(400, "No readable text found at that URL")
    return title, text


# ── Chunking ───────────────────────────────────────────────────────────────────


def chunk_text(text: str, chunk_size: int = 350, overlap: int = 60) -> list[str]:
    """
    Split into overlapping word windows.

    The overlap means a sentence straddling a boundary still appears whole in one
    of the two chunks, so retrieval does not lose it.

      chunk_size=350, overlap=60
        chunk 1: words 0-350
        chunk 2: words 290-640
        chunk 3: words 580-930
    """
    if overlap >= chunk_size:
        raise ValueError("overlap must be smaller than chunk_size")

    words = text.split()
    if not words:
        return []

    chunks: list[str] = []
    step = chunk_size - overlap
    for start in range(0, len(words), step):
        chunk = " ".join(words[start : start + chunk_size])
        if chunk.strip():
            chunks.append(chunk)
        if start + chunk_size >= len(words):
            break
    return chunks


def extract_and_chunk(file_path: str, file_type: str) -> list[str]:
    if file_type == "pdf":
        text = extract_text_from_pdf(file_path)
    elif file_type == "docx":
        text = extract_text_from_docx(file_path)
    elif file_type == "txt":
        text = extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    return chunk_text(text)


# ── Spreadsheets ───────────────────────────────────────────────────────────────


def get_dataframe(file_path: str, file_type: str) -> pd.DataFrame:
    if file_type == "csv":
        return pd.read_csv(file_path)
    if file_type == "xlsx":
        return pd.read_excel(file_path)
    raise ValueError(f"Not a spreadsheet: {file_type}")


def describe_dataframe(df: pd.DataFrame) -> str:
    """
    A compact schema description for the model.

    This is what lets the agent write correct pandas code without ever seeing the
    data: it gets column names, dtypes, and a couple of sample values, which is
    enough to reason about — and is a few hundred tokens instead of the entire
    dataset.
    """
    lines = [f"{len(df)} rows x {len(df.columns)} columns", "", "Columns:"]
    for col in df.columns:
        samples = df[col].dropna().unique()[:3]
        sample_text = ", ".join(str(s)[:40] for s in samples)
        lines.append(f"  - {col} ({df[col].dtype}) e.g. {sample_text}")
    return "\n".join(lines)
